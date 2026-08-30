"""
generate_survey_docx.py - Create a survey Word document from generated stimuli
================================================================================
Reads survey_stimuli.csv and the corresponding images, then builds
a .docx with one stimulus per page: floor plan + bubble diagram + rating.

Participants can fill in responses digitally.

Requirements:
    pip install python-docx

Usage:
    python generate_survey_docx.py
    python generate_survey_docx.py --dir survey_stimuli/ --out survey.docx
"""

from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path

try:
    _SCRIPT_DIR = Path(__file__).resolve().parent
except NameError:
    _SCRIPT_DIR = Path("/content/drive/MyDrive/bubble_diagram_project")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Likert scale questions ────────────────────────────────────────────────────
QUESTIONS = [
    "Q1. The bubble diagram accurately represents the rooms in the floor plan.",
    "Q2. The room adjacencies (connections) are correctly captured.",
    "Q3. The room sizes in the diagram are proportional to the actual plan.",
    "Q4. The diagram is easy to read and interpret.",
    "Q5. Overall, I would rate this bubble diagram as useful for understanding the layout.",
]

SCALE_OPTIONS = ["1\nStrongly\nDisagree", "2\nDisagree", "3\nNeutral", "4\nAgree", "5\nStrongly\nAgree"]


def _set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _add_styled_paragraph(doc, text, size=11, bold=False, color=None,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                          space_before=0):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _build_cover(doc):
    """Build the cover/instructions page."""
    # title
    _add_styled_paragraph(doc, "", size=11, space_after=60)
    _add_styled_paragraph(doc, "Floor Plan Bubble Diagram",
                          size=24, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_styled_paragraph(doc, "User Evaluation Survey",
                          size=18, color=(0x44, 0x44, 0x44),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # divider
    _add_styled_paragraph(doc, "_" * 72,
                          size=8, color=(0xCC, 0xCC, 0xCC),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # instructions
    _add_styled_paragraph(doc, "Instructions",
                          size=14, bold=True, space_after=12)

    instructions = [
        "This survey evaluates the quality of automatically generated bubble "
        "diagrams from architectural floor plan images.",
        "",
        "Each page presents:",
        "     \u2022  Left: The original floor plan image",
        "     \u2022  Right: The generated bubble diagram",
        "",
        "For each stimulus, please rate the bubble diagram on a scale of "
        "1 (Strongly Disagree) to 5 (Strongly Agree) for each question.",
        "",
        "There are no right or wrong answers. Please provide your honest "
        "assessment based on your understanding of the floor plan.",
        "",
        "Thank you for your participation.",
    ]

    for line in instructions:
        _add_styled_paragraph(doc, line, size=11, space_after=2)

    _add_styled_paragraph(doc, "", space_after=30)

    # participant info
    _add_styled_paragraph(doc, "_" * 72,
                          size=8, color=(0xCC, 0xCC, 0xCC),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=15)

    _add_styled_paragraph(doc, "Participant Information",
                          size=14, bold=True, space_after=15)

    # info table (clean fillable fields)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    fields = [
        ("Participant ID:", ""),
        ("Date:", ""),
        ("Background:", "[ ] Architect    [ ] Engineer    [ ] Student    [ ] Other: ________"),
        ("Years of experience:", ""),
    ]

    for i, (label, value) in enumerate(fields):
        label_cell = info_table.cell(i, 0)
        value_cell = info_table.cell(i, 1)

        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(11)
        lr.font.name = "Times New Roman"
        lr.bold = True
        lp.paragraph_format.space_after = Pt(8)

        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value if value else "___________________________")
        vr.font.size = Pt(11)
        vr.font.name = "Times New Roman"
        vp.paragraph_format.space_after = Pt(8)

    # page break
    doc.add_page_break()


def _build_stimulus_page(doc, stim_id, fp_path, bubble_path, is_last=False):
    """Build one stimulus page."""
    # ── header ────────────────────────────────────────────────────────────────
    _add_styled_paragraph(doc, f"Stimulus {stim_id}",
                          size=14, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_styled_paragraph(doc, "_" * 72,
                          size=8, color=(0xCC, 0xCC, 0xCC),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ── images in a 1x2 table ────────────────────────────────────────────────
    img_table = doc.add_table(rows=2, cols=2)
    img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # titles row
    for j, title in enumerate(["Floor Plan", "Bubble Diagram"]):
        cell = img_table.cell(0, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        r.bold = True
        p.paragraph_format.space_after = Pt(4)

    # images row
    for j, img_path in enumerate([fp_path, bubble_path]):
        cell = img_table.cell(1, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path.exists():
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(2.8))
        else:
            p.add_run("(image not found)")
        p.paragraph_format.space_after = Pt(4)

    _add_styled_paragraph(doc, "", space_after=6)

    # ── divider ──────────────────────────────────────────────────────────────
    _add_styled_paragraph(doc, "_" * 72,
                          size=8, color=(0xCC, 0xCC, 0xCC),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # ── scale header ─────────────────────────────────────────────────────────
    _add_styled_paragraph(
        doc,
        "1 = Strongly Disagree    2 = Disagree    3 = Neutral    4 = Agree    5 = Strongly Agree",
        size=8, color=(0x66, 0x66, 0x66),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
    )

    # ── rating table ─────────────────────────────────────────────────────────
    n_q = len(QUESTIONS)
    rating_table = doc.add_table(rows=n_q + 1, cols=6)
    rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    headers = ["Question"] + SCALE_OPTIONS
    for j, h in enumerate(headers):
        cell = rating_table.cell(0, j)
        _set_cell_shading(cell, "4472C4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(7 if j > 0 else 9)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.bold = True
        p.paragraph_format.space_after = Pt(2)

    # question rows
    for i, q in enumerate(QUESTIONS):
        row_idx = i + 1
        bg = "FFFFFF" if row_idx % 2 == 1 else "D9E2F3"

        # question cell
        q_cell = rating_table.cell(row_idx, 0)
        _set_cell_shading(q_cell, bg)
        qp = q_cell.paragraphs[0]
        qr = qp.add_run(q)
        qr.font.size = Pt(9)
        qr.font.name = "Times New Roman"
        qp.paragraph_format.space_after = Pt(4)

        # rating cells (empty circles for filling in)
        for j in range(1, 6):
            r_cell = rating_table.cell(row_idx, j)
            _set_cell_shading(r_cell, bg)
            rp = r_cell.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = rp.add_run("\u25cb")  # empty circle
            rr.font.size = Pt(14)
            rr.font.name = "Times New Roman"
            rp.paragraph_format.space_after = Pt(2)

    # set column widths
    for row in rating_table.rows:
        row.cells[0].width = Cm(10)
        for j in range(1, 6):
            row.cells[j].width = Cm(1.6)

    _add_styled_paragraph(doc, "", space_after=8)

    # ── comments ─────────────────────────────────────────────────────────────
    _add_styled_paragraph(doc, "Comments (optional):",
                          size=10, bold=True, space_after=4)

    # blank lines for comments
    for _ in range(3):
        _add_styled_paragraph(doc, "_" * 80,
                              size=9, color=(0xDD, 0xDD, 0xDD), space_after=8)

    # page break (except after last page)
    if not is_last:
        doc.add_page_break()


def generate_docx(stim_dir, out_path):
    stim_dir = Path(stim_dir)
    csv_path = stim_dir / "survey_stimuli.csv"

    if not csv_path.exists():
        raise FileNotFoundError(f"CSV not found: {csv_path}")

    stimuli = []
    with open(csv_path) as f:
        for row in csv.DictReader(f):
            stimuli.append(row)

    print(f"[survey] {len(stimuli)} stimuli from {csv_path}")

    doc = Document()

    # set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # set A4 margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # cover page
    _build_cover(doc)

    # stimulus pages
    for i, s in enumerate(stimuli):
        fp_path = stim_dir / "floorplans" / s["Floor_Plan_File"]
        bubble_path = stim_dir / "bubbles" / s["Bubble_Diagram_File"]
        is_last = (i == len(stimuli) - 1)
        _build_stimulus_page(doc, s["Stimulus_ID"], fp_path, bubble_path, is_last)

    doc.save(out_path)
    print(f"Saved -> {out_path}  ({len(stimuli) + 1} pages)")


def main():
    parser = argparse.ArgumentParser(
        description="Generate survey DOCX from stimuli images")
    parser.add_argument("--dir", type=str,
                        default=str(_SCRIPT_DIR / "survey_stimuli"),
                        help="Directory with floorplans/, bubbles/, survey_stimuli.csv")
    parser.add_argument("--out", type=str,
                        default=str(_SCRIPT_DIR / "survey.docx"),
                        help="Output DOCX path")
    args = parser.parse_args()

    generate_docx(args.dir, args.out)


if __name__ == "__main__":
    main()
